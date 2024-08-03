import os
import re
from docx import Document
from Managers.filemanager import FileManager

class WordFileManager(FileManager):
    def __init__(self, path_file: str) -> None:
        super().__init__(path_file)
        self.document = Document(self.path_file)
        self.extension = 'docx'
        self.set_file_name()
    
    def split_file(self) -> None:
        output_dir = f"{self.file_name}"
        os.makedirs(output_dir, exist_ok=True)
        
        for i, paragraph in enumerate(self.document.paragraphs):
            new_doc = Document()
            new_doc.add_paragraph(paragraph.text)
            
            output_docx = f"{output_dir}/page_{i + 1}.{self.extension}"
            new_doc.save(output_docx)

    def merge_files(self, output_docx: str) -> None:
        merged_doc = Document()
        input_dir = f"{self.file_name}"
        page_files = [f"{input_dir}/{file}" for file in os.listdir(input_dir) if file.endswith('.docx')]
        page_files.sort(key=lambda f: int(re.search(r'page_(\d+).docx', f).group(1)))

        for file in page_files:
            doc = Document(file)
            for paragraph in doc.paragraphs:
                merged_doc.add_paragraph(paragraph.text)
        
        merged_doc.save(f'{output_docx}_merged.{self.extension}')
        print(f"Word files merged into {output_docx}_merged.{self.extension}")

